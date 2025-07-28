from docx import Document
from docx.shared import RGBColor
import json

def rgb_from_color(color):
    if color and hasattr(color, 'rgb'):
        return color.rgb.__str__()
    return "None"

def extract_elements(doc_path):
    doc = Document(doc_path)
    elements = []

    for para in doc.paragraphs:
        el = {
            'text': para.text.strip(),
            'style': para.style.name,
            'font_size': None,
            'font_color': None,
            'background_color': None,
            'type': 'paragraph'
        }
        if para.runs:
            run = para.runs[0]
            if run.font.size:
                el['font_size'] = str(run.font.size.pt)
            el['font_color'] = rgb_from_color(run.font.color)
        elements.append(el)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                elements.append({
                    'text': cell.text.strip(),
                    'style': cell.paragraphs[0].style.name,
                    'type': 'table_cell'
                })
    return elements
